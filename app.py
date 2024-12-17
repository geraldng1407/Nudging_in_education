
# '''
# docker build -t nudging_in_education .
# docker run -p 8501:8501 --env-file ./.env nudging_in_education
# '''

# import streamlit as st
# from classes.chunker import Chunker
# from classes.content_generator import ContentGenerator
# import streamlit.components.v1 as components
# from prompts.prompts import MULTI_CHOICE_PROMPT
# from classes.output_parser import ScenariosWithQuestions, ScenarioWithQuestions, Question
# import os
# from datetime import datetime
# import json
# import threading
# from dotenv import load_dotenv
# import pandas as pd

# load_dotenv()
# feedback_lock = threading.Lock()

# # Function to create Mermaid code (unchanged)
# def create_mermaid_code(code):
#     return f"""
#     <html>
#     <head>
#     <style>
#     .mermaid {{
#         display: grid;
#         place-items: center;
#         height: 100%;
#     }}
#     .mermaid svg {{
#         max-width: 100%;
#     }}
#     </style>
#     </head>
#     <body>
#     <pre class="mermaid">
#     {code}
#     </pre>
#     <script type="module">
#     import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
#     mermaid.initialize({{ startOnLoad: true }});
#     </script>
#     </body>
#     </html>
#     """

# def get_feedback_data():
#     """
#     Read feedback data from feedback.json and return as a pandas DataFrame.
#     """
#     feedback_file = "feedback.json"
#     feedback_path = os.path.join(os.getcwd(), feedback_file)
#     valid_study_types = ["Visual", "Summary", "Scenario"]

#     if os.path.exists(feedback_path):
#         with feedback_lock:
#             with open(feedback_path, 'r', encoding='utf-8') as f:
#                 data = json.load(f)

#         feedback_list = []
#         for study_type in valid_study_types:
#             feedbacks = data.get(study_type, [])
#             for fb in feedbacks:
#                 fb_copy = fb.copy()
#                 fb_copy['study_type'] = study_type
#                 feedback_list.append(fb_copy)

#         if feedback_list:
#             df = pd.DataFrame(feedback_list)
#             return df
#         else:
#             return pd.DataFrame()
#     else:
#         return pd.DataFrame()

# def main():
#     # Set page configuration
#     st.set_page_config(layout="wide")

#     # Initialize session state variables
#     if 'feedback_open' not in st.session_state:
#         st.session_state.feedback_open = False

#     # Apply custom CSS for the floating button
#     st.markdown("""
#     <style>
#     .floating-button {
#         position: fixed;  
#         bottom: 30px;
#         right: 30px;
#         background-color: #47b3b3;
#         color: white;
#         border: none;
#         border-radius: 50%;
#         width: 70px;
#         height: 70px;
#         text-align: center;
#         font-size: 32px;
#         box-shadow: 0 4px 8px rgba(0,0,0,0.3);
#         cursor: pointer;
#         z-index: 1000;  /* Ensures button is always on top */
#         transition: background-color 0.3s, transform 0.3s;
#     }
#     .floating-button:hover {
#         background-color: #2e8b8b;
#         transform: scale(1.1);
#     }
#     </style>
#     """, unsafe_allow_html=True)
    
#     # Floating feedback button with toggle functionality
#     if st.button("💬Feedback", key="feedback_button"):
#         st.session_state.feedback_open = not st.session_state.feedback_open

#     # Centered upload button and role toggle
#     with st.container():
#         col1, col2, col3 = st.columns([1, 4, 1])

#         with col3:
#             role = st.radio("Mode", ["User", "Admin"], index=0, key="role_toggle", horizontal=True)

#         with col2:
#             uploaded_file = st.file_uploader("Upload", type=["txt", "docx"])
#             if uploaded_file is not None:
#                 if 'chunked_sections' not in st.session_state:
#                     # Read the file as a string and process only once
#                     if uploaded_file.type == "text/plain":
#                         doc = uploaded_file.read().decode("utf-8")
#                     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                         import docx
#                         docx_document = docx.Document(uploaded_file)
#                         doc = "\n\n".join(
#                             [para.text for para in docx_document.paragraphs])
#                     else:
#                         st.error("Unsupported file type.")
#                         doc = ""

#                     if doc:
#                         chunker = Chunker()
#                         with col2:
#                             with st.spinner('Uploading File...'):
#                                 chunked_sections = chunker.split(doc)
#                                 section_titles = list(chunked_sections.keys())
#                                 section_bodies = list(
#                                     chunked_sections.values())
#                                 st.toast(
#                                     'File uploaded successfully!', icon="✅")

#                         # Store results in session state so they persist across reruns
#                         st.session_state.chunked_sections = chunked_sections
#                         st.session_state.section_titles = section_titles
#                         st.session_state.section_bodies = section_bodies
#             else:
#                 # Resetting session state if no file is uploaded or new file is selected
#                 if 'chunked_sections' in st.session_state:
#                     del st.session_state['chunked_sections']
#                     del st.session_state['section_titles']
#                     del st.session_state['section_bodies']

#     # Display the feedback form in the sidebar if feedback_open is True
#     if st.session_state.feedback_open:
#         with st.sidebar:
#             st.header("Feedback")
#             with st.form("feedback_form"):
#                 feedback_rating = st.select_slider(
#                     "Rate the quality of the output:",
#                     options=["1 - Very Poor", "2 - Poor", "3 - Average", "4 - Good", "5 - Excellent"]
#                 )
#                 feedback_comments = st.text_area("Additional Comments:")
#                 submitted = st.form_submit_button("Submit Feedback")
                
#                 if submitted:
#                     # Retrieve the current study_type from session_state
#                     study_type = st.session_state.get('current_study_type', None)
                    
#                     if study_type is None:
#                         st.error("Please select a study type before submitting feedback.")
#                     else:
#                         feedback = {
#                             "rating": feedback_rating,
#                             "comments": feedback_comments,
#                         }
#                         save_feedback(feedback, study_type)
#                         st.success("Thank you for your feedback!")
                        
#                         # Immediately close the feedback sidebar
#                         st.session_state.feedback_open = False

#     if 'section_titles' in st.session_state:
#         with col2:
#             selected_title = st.selectbox(
#                 "Choose a section:",
#                 options=st.session_state.section_titles,
#                 index=None,
#                 placeholder="Select a section",
#             )
#             st.markdown("<hr style='border:1px solid #f0f0f0'>",
#                         unsafe_allow_html=True)

#             if selected_title is not None:
#                 st.header("Study Type")
#                 study_type = st.radio(
#                     "Study Type",
#                     ["Visual", "Summary", "Scenario"],  # Unified 'Scenario'
#                     horizontal=True,
#                     index=0,
#                     label_visibility="collapsed"
#                 )
#                 st.session_state.current_study_type = study_type
#                 content_generator = ContentGenerator()

#                 if "generated_content" not in st.session_state:
#                     st.session_state.generated_content = {}

#                 key = (selected_title, study_type)
                
#                 # Retrieve feedback for this study type from session state
#                 feedback_for_study_type = st.session_state.feedback.get(study_type, []) if 'feedback' in st.session_state else []

#                 if study_type == "Visual":
#                     regenerate = st.button("Regenerate Mindmap")
#                     if regenerate or key not in st.session_state.generated_content:
#                         with st.spinner("Generating Mindmap..."):
#                             try:
#                                 mermaid_code = content_generator.create_visual(
#                                     st.session_state.section_bodies[st.session_state.section_titles.index(
#                                         selected_title)],
#                                     feedback=feedback_for_study_type  # Pass feedback here
#                                 )
#                                 st.session_state.generated_content[key] = mermaid_code
#                             except ValueError as ve:
#                                 st.error(f"Error generating visual: {ve}")

#                     if key in st.session_state.generated_content:
#                         components.html(create_mermaid_code(
#                             st.session_state.generated_content[key]), height=1000)

#                 elif study_type == "Summary":
#                     regenerate = st.button("Regenerate Summary")
#                     if regenerate or key not in st.session_state.generated_content:
#                         with st.spinner("Generating Summary..."):
#                             try:
#                                 summary = content_generator.create_summary(
#                                     st.session_state.section_bodies[st.session_state.section_titles.index(
#                                         selected_title)],
#                                     feedback=feedback_for_study_type  # Pass feedback here
#                                 )
#                                 st.session_state.generated_content[key] = summary
#                             except ValueError as ve:
#                                 st.error(f"Error generating summary: {ve}")

#                     if key in st.session_state.generated_content:
#                         summary_content = st.session_state.generated_content[key]
#                         st.markdown(summary_content, unsafe_allow_html=True)

#                 elif study_type == "Scenario":
#                     regenerate = st.button("Regenerate Scenario and Questions")
#                     if regenerate or key not in st.session_state.generated_content:
#                         with st.spinner("Generating Scenarios and Questions..."):
#                             try:
#                                 content_text = st.session_state.section_bodies[st.session_state.section_titles.index(
#                                     selected_title)]
#                                 scenarios = content_generator.create_scenarios_with_questions(
#                                     content_text,
#                                     feedback=feedback_for_study_type  # Pass feedback here
#                                 )
#                                 st.session_state.generated_content[key] = scenarios
#                             except ValueError as ve:
#                                 st.error(
#                                     f"Error generating scenarios and questions: {ve}")

#                     if key in st.session_state.generated_content:
#                         scenarios = st.session_state.generated_content[key]
#                         for idx, scenario in enumerate(scenarios, start=1):
#                             st.subheader(f"Scenario {idx}")
#                             st.markdown(
#                                 f"\n\n{scenario.get('scenario')}", unsafe_allow_html=True)

#                             st.markdown("**Questions:**")
#                             user_answers = {}
#                             for q_idx, question in enumerate(scenario.get("questions"), start=1):
#                                 question_number = f"Q{idx}.{q_idx}"

#                                 # Define option letters
#                                 option_letters = ['A', 'B', 'C', 'D']

#                                 # Ensure there are exactly four options
#                                 if len(question.get("options")) != 4:
#                                     st.error(
#                                         f"Question {question_number} does not have exactly four options.")
#                                     continue

#                                 # Map letters to options
#                                 options_with_letters = {
#                                     letter: option.strip() for letter, option in zip(option_letters, question.get("options"))
#                                 }

#                                 # Create a list of labels like "Option Text"
#                                 option_labels = [
#                                     f"{text}" for letter, text in options_with_letters.items()]
                                
#                                 # Use the option letters as the actual values
#                                 selected_letter = st.radio(
#                                     f"**{question_number}: {question.get('question')}**", 
#                                     options=option_letters,
#                                     format_func=lambda x: option_labels[option_letters.index(x)], 
#                                     key=f"{key}_{idx}_{q_idx}",
#                                     index=None
#                                 )

#                                 user_answers[q_idx] = selected_letter

#                             if st.button(f"Submit Answers for Scenario {idx}"):
#                                 st.subheader(f"Results for Scenario {idx}")
#                                 for q_idx, question in enumerate(scenario.get("questions"), start=1):
#                                     selected = user_answers.get(q_idx, None)
#                                     correct = question.get("answer")
#                                     if selected == correct:
#                                         st.success(
#                                             f"Question {q_idx}: Correct!")
#                                     else:
#                                         # Find the correct option text for display
#                                         try:
#                                             correct_text = next(
#                                                 option for letter, option in zip(option_letters, question.get("options"))
#                                                 if letter == correct
#                                             )
#                                             st.error(f"Question {q_idx}: Incorrect.\nCorrect Answer: {correct_text}")

#                                         except StopIteration:
#                                             st.error(
#                                                 f"Question {q_idx}: Incorrect.  \n Correct Answer: {correct}")

#                             st.markdown("---")  # Separator between scenarios

#     # Show admin functions only when in Admin mode
#     if role == "Admin":
#         st.markdown("---")  # Optional: Add a divider line

#         st.header("Admin Functions")
#         st.write("Download all feedbacks as a CSV file.")

#         feedback_df = get_feedback_data()
#         if not feedback_df.empty:
#             # Convert DataFrame to CSV
#             csv_data = feedback_df.to_csv(index=False).encode('utf-8')

#             # Provide the download button
#             st.download_button(
#                 label="Download Feedback CSV",
#                 data=csv_data,
#                 file_name='feedbacks.csv',
#                 mime='text/csv',
#                 key="download_csv_button"
#             )
#         else:
#             st.info("No feedback data available to download.")

# def save_feedback(feedback, study_type):
#     """
#     Save feedback to a JSON file categorized by study_type.

#     The JSON structure will be:
#     {
#         "Visual": [ ... ],
#         "Summary": [ ... ],
#         "Scenario": [ ... ]
#     }
#     """
#     feedback_file = "feedback.json"
#     feedback_path = os.path.join(os.getcwd(), feedback_file)
#     valid_study_types = ["Visual", "Summary", "Scenario"]

#     if study_type not in valid_study_types:
#         st.error(f"Invalid study type: {study_type}")
#         return

#     try:
#         with feedback_lock:
#             if os.path.exists(feedback_path):
#                 with open(feedback_path, 'r', encoding='utf-8') as f:
#                     data = json.load(f)
#             else:
#                 # Initialize the JSON structure if it doesn't exist
#                 data = {stype: [] for stype in valid_study_types}

#             # Ensure all study_type keys exist
#             for stype in valid_study_types:
#                 if stype not in data:
#                     data[stype] = []

#             # Add datetime to feedback
#             feedback["datetime"] = datetime.now().isoformat()

#             # Append the feedback to the appropriate study_type list
#             data[study_type].append(feedback)

#             # Write the updated data back to the JSON file
#             with open(feedback_path, 'w', encoding='utf-8') as f:
#                 json.dump(data, f, ensure_ascii=False, indent=4)

#             # Save feedback into session state
#             if 'feedback' not in st.session_state:
#                 st.session_state.feedback = {stype: [] for stype in valid_study_types}
#             st.session_state.feedback[study_type].append(feedback)

#     except Exception as e:
#         st.error(f"Failed to save feedback: {e}")

# if __name__ == "__main__":
#     main()

import streamlit as st
from classes.chunker import Chunker
from classes.content_generator import ContentGenerator
import streamlit.components.v1 as components
from prompts.prompts import MULTI_CHOICE_PROMPT
from classes.output_parser import ScenariosWithQuestions, ScenarioWithQuestions, Question
import os
from datetime import datetime
import json
import threading
from dotenv import load_dotenv
import pandas as pd

load_dotenv()
feedback_lock = threading.Lock()

# Function to create Mermaid code (unchanged)
def create_mermaid_code(code):
    return f"""
    <html>
    <head>
    <style>
    .mermaid {{
        display: grid;
        place-items: center;
        height: 100%;
    }}
    .mermaid svg {{
        max-width: 100%;
    }}
    </style>
    </head>
    <body>
    <pre class="mermaid">
    {code}
    </pre>
    <script type="module">
    import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
    mermaid.initialize({{ startOnLoad: true }});
    </script>
    </body>
    </html>
    """

def get_feedback_data():
    """
    Read feedback data from feedback.json and return as a pandas DataFrame.
    """
    feedback_file = "feedback.json"
    feedback_path = os.path.join(os.getcwd(), feedback_file)
    valid_study_types = ["Visual", "Summary", "Scenario"]

    if os.path.exists(feedback_path):
        with feedback_lock:
            with open(feedback_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

        feedback_list = []
        for study_type in valid_study_types:
            feedbacks = data.get(study_type, [])
            for fb in feedbacks:
                fb_copy = fb.copy()
                fb_copy['study_type'] = study_type
                feedback_list.append(fb_copy)

        if feedback_list:
            df = pd.DataFrame(feedback_list)
            return df
        else:
            return pd.DataFrame()
    else:
        return pd.DataFrame()

def main():
    # Set page configuration
    st.set_page_config(layout="wide")

    # Initialize session state variables
    if 'feedback_open' not in st.session_state:
        st.session_state.feedback_open = False

    # Apply custom CSS for the floating button
    st.markdown("""
    <style>
    .floating-button {
        position: fixed;  
        bottom: 30px;
        right: 30px;
        background-color: #47b3b3;
        color: white;
        border: none;
        border-radius: 50%;
        width: 70px;
        height: 70px;
        text-align: center;
        font-size: 32px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.3);
        cursor: pointer;
        z-index: 1000;  /* Ensures button is always on top */
        transition: background-color 0.3s, transform 0.3s;
    }
    .floating-button:hover {
        background-color: #2e8b8b;
        transform: scale(1.1);
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Floating feedback button with toggle functionality
    if st.button("💬Feedback", key="feedback_button"):
        st.session_state.feedback_open = not st.session_state.feedback_open

    # Centered upload button and role toggle
    with st.container():
        col1, col2, col3 = st.columns([1, 4, 1])

        with col3:
            role = st.radio("Mode", ["User", "Admin"], index=0, key="role_toggle", horizontal=True)

        with col2:
            uploaded_file = st.file_uploader("Upload", type=["txt", "docx"])
            if uploaded_file is not None:
                if 'chunked_sections' not in st.session_state:
                    # Read the file as a string and process only once
                    if uploaded_file.type == "text/plain":
                        doc = uploaded_file.read().decode("utf-8")
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        import docx
                        docx_document = docx.Document(uploaded_file)
                        doc = "\n\n".join(
                            [para.text for para in docx_document.paragraphs])
                    else:
                        st.error("Unsupported file type.")
                        doc = ""

                    if doc:
                        chunker = Chunker()
                        with col2:
                            with st.spinner('Uploading File...'):
                                chunked_sections = chunker.split(doc)
                                section_titles = list(chunked_sections.keys())
                                section_bodies = list(
                                    chunked_sections.values())
                                st.toast(
                                    'File uploaded successfully!', icon="✅")

                        # Store results in session state so they persist across reruns
                        st.session_state.chunked_sections = chunked_sections
                        st.session_state.section_titles = section_titles
                        st.session_state.section_bodies = section_bodies
            else:
                # Resetting session state if no file is uploaded or new file is selected
                if 'chunked_sections' in st.session_state:
                    del st.session_state['chunked_sections']
                    del st.session_state['section_titles']
                    del st.session_state['section_bodies']

    # Display the feedback form in the sidebar if feedback_open is True
    if st.session_state.feedback_open:
        with st.sidebar:
            st.header("Feedback")
            with st.form("feedback_form"):
                feedback_rating = st.select_slider(
                    "Rate the quality of the output:",
                    options=["1 - Very Poor", "2 - Poor", "3 - Average", "4 - Good", "5 - Excellent"]
                )
                feedback_comments = st.text_area("Additional Comments:")
                submitted = st.form_submit_button("Submit Feedback")
                
                if submitted:
                    # Retrieve the current study_type from session_state
                    study_type = st.session_state.get('current_study_type', None)
                    
                    if study_type is None:
                        st.error("Please select a study type before submitting feedback.")
                    else:
                        feedback = {
                            "rating": feedback_rating,
                            "comments": feedback_comments,
                        }
                        save_feedback(feedback, study_type)
                        st.success("Thank you for your feedback!")
                        
                        # Immediately close the feedback sidebar
                        st.session_state.feedback_open = False

    if 'section_titles' in st.session_state:
        with col2:
            selected_title = st.selectbox(
                "Choose a section:",
                options=st.session_state.section_titles,
                index=0,
                placeholder="Select a section",
            )
            st.markdown("<hr style='border:1px solid #f0f0f0'>",
                        unsafe_allow_html=True)

            if selected_title is not None:
                st.header("Study Type")
                study_type = st.radio(
                    "Study Type",
                    ["Visual", "Summary", "Scenario"],  # Unified 'Scenario'
                    horizontal=True,
                    index=0,
                    label_visibility="collapsed"
                )
                st.session_state.current_study_type = study_type
                content_generator = ContentGenerator()

                if "generated_content" not in st.session_state:
                    st.session_state.generated_content = {}

                key = (selected_title, study_type)
                
                # Retrieve feedback for this study type from session state
                feedback_for_study_type = st.session_state.feedback.get(study_type, []) if 'feedback' in st.session_state else []

                if study_type == "Visual":
                    regenerate = st.button("Regenerate Mindmap")
                    if regenerate or key not in st.session_state.generated_content:
                        with st.spinner("Generating Mindmap..."):
                            try:
                                mermaid_code = content_generator.create_visual(
                                    st.session_state.section_bodies[st.session_state.section_titles.index(
                                        selected_title)],
                                    feedback=feedback_for_study_type  # Pass feedback here
                                )
                                st.session_state.generated_content[key] = mermaid_code
                            except ValueError as ve:
                                st.error(f"Error generating visual: {ve}")

                    if key in st.session_state.generated_content:
                        components.html(create_mermaid_code(
                            st.session_state.generated_content[key]), height=1000)

                elif study_type == "Summary":
                    regenerate = st.button("Regenerate Summary")
                    if regenerate or key not in st.session_state.generated_content:
                        with st.spinner("Generating Summary..."):
                            try:
                                summary = content_generator.create_summary(
                                    st.session_state.section_bodies[st.session_state.section_titles.index(
                                        selected_title)],
                                    feedback=feedback_for_study_type  # Pass feedback here
                                )
                                st.session_state.generated_content[key] = summary
                            except ValueError as ve:
                                st.error(f"Error generating summary: {ve}")

                    if key in st.session_state.generated_content:
                        summary_content = st.session_state.generated_content[key]
                        st.markdown(summary_content, unsafe_allow_html=True)

                elif study_type == "Scenario":
                    regenerate = st.button("Regenerate Scenario and Questions")
                    if regenerate or key not in st.session_state.generated_content:
                        with st.spinner("Generating Scenarios and Questions..."):
                            try:
                                content_text = st.session_state.section_bodies[st.session_state.section_titles.index(
                                    selected_title)]
                                scenarios = content_generator.create_scenarios_with_questions(
                                    content_text,
                                    feedback=feedback_for_study_type  # Pass feedback here
                                )
                                st.session_state.generated_content[key] = scenarios
                            except ValueError as ve:
                                st.error(
                                    f"Error generating scenarios and questions: {ve}")

                    if key in st.session_state.generated_content:
                        scenarios = st.session_state.generated_content[key]
                        for idx, scenario in enumerate(scenarios, start=1):
                            st.subheader(f"Scenario {idx}")
                            st.markdown(
                                f"\n\n{scenario.get('scenario')}", unsafe_allow_html=True)

                            st.markdown("**Questions:**")
                            user_answers = {}
                            for q_idx, question in enumerate(scenario.get("questions"), start=1):
                                question_number = f"Q{idx}.{q_idx}"

                                # Define option letters
                                option_letters = ['A', 'B', 'C', 'D']

                                # Ensure there are exactly four options
                                if len(question.get("options")) != 4:
                                    st.error(
                                        f"Question {question_number} does not have exactly four options.")
                                    continue

                                # Map letters to options
                                options_with_letters = {
                                    letter: option.strip() for letter, option in zip(option_letters, question.get("options"))
                                }

                                # Create a list of labels like "Option Text"
                                option_labels = [
                                    f"{text}" for letter, text in options_with_letters.items()]
                                
                                # Use the option letters as the actual values
                                selected_letter = st.radio(
                                    f"**{question_number}: {question.get('question')}**", 
                                    options=option_letters,
                                    format_func=lambda x: option_labels[option_letters.index(x)], 
                                    key=f"{key}_{idx}_{q_idx}",
                                    index=None
                                )

                                user_answers[q_idx] = selected_letter

                            if st.button(f"Submit Answers for Scenario {idx}"):
                                st.subheader(f"Results for Scenario {idx}")
                                for q_idx, question in enumerate(scenario.get("questions"), start=1):
                                    selected = user_answers.get(q_idx, None)
                                    correct = question.get("answer")
                                    if selected == correct:
                                        st.success(
                                            f"Question {q_idx}: Correct!")
                                    else:
                                        # Find the correct option text for display
                                        try:
                                            correct_text = next(
                                                option for letter, option in zip(option_letters, question.get("options"))
                                                if letter == correct
                                            )
                                            st.error(f"Question {q_idx}: Incorrect.\nCorrect Answer: {correct_text}")

                                        except StopIteration:
                                            st.error(
                                                f"Question {q_idx}: Incorrect.  \n Correct Answer: {correct}")

                            st.markdown("---")  # Separator between scenarios

    # Show admin functions only when in Admin mode
    if role == "Admin":
        st.markdown("---")  # Optional: Add a divider line

        st.header("Admin Functions")
        st.write("Download all feedbacks as a CSV file.")

        feedback_df = get_feedback_data()
        if not feedback_df.empty:
            # Convert DataFrame to CSV
            csv_data = feedback_df.to_csv(index=False).encode('utf-8')

            # Provide the download button
            st.download_button(
                label="Download Feedback CSV",
                data=csv_data,
                file_name='feedbacks.csv',
                mime='text/csv',
                key="download_csv_button"
            )
        else:
            st.info("No feedback data available to download.")

        # New button to view feedback in session state
        st.write("View all feedbacks in the current session.")
        
        if 'view_feedback_open' not in st.session_state:
            st.session_state.view_feedback_open = False

        if st.button("View Feedback in Session", key="view_feedback_button"):
            st.session_state.view_feedback_open = not st.session_state.view_feedback_open

        if st.session_state.view_feedback_open:
            if 'feedback' in st.session_state and st.session_state.feedback:
                st.subheader("Feedback in Current Session")
                feedback_list = []
                for study_type, feedbacks in st.session_state.feedback.items():
                    for fb in feedbacks:
                        fb_copy = fb.copy()
                        fb_copy['study_type'] = study_type
                        feedback_list.append(fb_copy)
                if feedback_list:
                    df_session = pd.DataFrame(feedback_list)
                    st.dataframe(df_session)
                else:
                    st.info("No feedback data in session state.")
            else:
                st.info("No feedback data in session state.")

def save_feedback(feedback, study_type):
    """
    Save feedback to a JSON file categorized by study_type.

    The JSON structure will be:
    {
        "Visual": [ ... ],
        "Summary": [ ... ],
        "Scenario": [ ... ]
    }
    """
    feedback_file = "feedback.json"
    feedback_path = os.path.join(os.getcwd(), feedback_file)
    valid_study_types = ["Visual", "Summary", "Scenario"]

    if study_type not in valid_study_types:
        st.error(f"Invalid study type: {study_type}")
        return

    try:
        with feedback_lock:
            if os.path.exists(feedback_path):
                with open(feedback_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            else:
                # Initialize the JSON structure if it doesn't exist
                data = {stype: [] for stype in valid_study_types}

            # Ensure all study_type keys exist
            for stype in valid_study_types:
                if stype not in data:
                    data[stype] = []

            # Add datetime to feedback
            feedback["datetime"] = datetime.now().isoformat()

            # Append the feedback to the appropriate study_type list
            data[study_type].append(feedback)

            # Write the updated data back to the JSON file
            with open(feedback_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)

            # Save feedback into session state
            if 'feedback' not in st.session_state:
                st.session_state.feedback = {stype: [] for stype in valid_study_types}
            st.session_state.feedback[study_type].append(feedback)

    except Exception as e:
        st.error(f"Failed to save feedback: {e}")

if __name__ == "__main__":
    main()